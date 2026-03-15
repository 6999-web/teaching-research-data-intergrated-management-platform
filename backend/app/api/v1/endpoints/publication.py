"""
公示相关的API端点

包括发起公示、结果分发等功能
"""

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import Response
from sqlalchemy.orm import Session
from typing import List, Optional
from uuid import UUID
from datetime import datetime
import io
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.deps import get_db, require_evaluation_office, require_management_roles, get_current_user
from app.models.user import User
from app.models.publication import Publication
from app.models.self_evaluation import SelfEvaluation
from app.models.approval import Approval
from app.models.operation_log import OperationLog
from app.models.manual_score import ManualScore
from app.models.final_score import FinalScore
from app.models.ai_score import AIScore
from app.models.attachment import Attachment
from app.models.teaching_office import TeachingOffice
from app.schemas.publication import (
    PublishRequest,
    PublishResponse,
    PublicationDetail,
    DistributeRequest,
    DistributeResponse,
)
from app.services.insight_service import generate_insight_for_evaluation
import logging

logger = logging.getLogger(__name__)
router = APIRouter()


def _manual_score_total(scores_json) -> float:
    """从 ManualScore.scores JSON 计算总分"""
    if not scores_json or not isinstance(scores_json, list):
        return 0.0
    total = 0.0
    for item in scores_json:
        try:
            if isinstance(item, dict) and "score" in item:
                total += float(item["score"])
        except (TypeError, ValueError):
            continue
    return total


@router.get("/evaluations-for-publication")
def get_evaluations_for_publication(
    year: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office),
):
    """
    获取可以发起公示的评估列表。
    包含考评小组的评分信息和附件。
    状态包括: finalized, manually_scored, approved, published, distributed
    """
    query = (
        db.query(SelfEvaluation)
        .join(TeachingOffice, SelfEvaluation.teaching_office_id == TeachingOffice.id)
        .filter(
            SelfEvaluation.status.in_([
                "manually_scored", "ready_for_final", "finalized", "approved", "published", "distributed"
            ])
        )
    )
    if year:
        query = query.filter(SelfEvaluation.evaluation_year == year)
    evaluations = query.order_by(SelfEvaluation.submitted_at.desc()).all()

    result = []
    for ev in evaluations:
        office = ev.teaching_office
        office_name = office.name if office else ""

        # 手动评分信息
        manual_scores = (
            db.query(ManualScore)
            .filter(ManualScore.evaluation_id == ev.id)
            .order_by(ManualScore.submitted_at.desc())
            .all()
        )
        manual_score_list = []
        for ms in manual_scores:
            scores_list = ms.scores if isinstance(ms.scores, list) else []
            total = _manual_score_total(scores_list)
            manual_score_list.append({
                "reviewer_name": ms.reviewer_name,
                "reviewer_role": ms.reviewer_role,
                "total": total,
                "submitted_at": ms.submitted_at.isoformat() if ms.submitted_at else None,
            })

        # AI评分
        ai_score = db.query(AIScore).filter(AIScore.evaluation_id == ev.id).first()
        ai_score_val = float(ai_score.total_score) if ai_score else None

        # 最终得分
        final = db.query(FinalScore).filter(FinalScore.evaluation_id == ev.id).first()
        final_score_val = float(final.final_score) if final else None

        # 附件（考评小组端上传的）
        attachments = (
            db.query(Attachment)
            .filter(Attachment.evaluation_id == ev.id)
            .all()
        )
        attachment_list = [
            {
                "id": str(a.id),
                "file_name": a.file_name,
                "indicator": a.indicator,
                "classified_by": a.classified_by,
                "uploaded_at": a.uploaded_at.isoformat() if a.uploaded_at else None,
            }
            for a in attachments
        ]

        # 判断是否已公示（是否在 publication 中）
        pub = db.query(Publication).filter(
            Publication.evaluation_ids.contains([str(ev.id)])
        ).first()

        result.append({
            "id": str(ev.id),
            "teaching_office_id": str(ev.teaching_office_id),
            "teaching_office_name": office_name,
            "evaluation_year": ev.evaluation_year,
            "status": ev.status,
            "ai_score": ai_score_val,
            "manual_scores": manual_score_list,
            "final_score": final_score_val,
            "attachments": attachment_list,
            "submitted_at": ev.submitted_at.isoformat() if ev.submitted_at else None,
            "is_published": pub is not None,
            "is_distributed": pub.distributed_at is not None if pub else False,
        })

    return result


@router.post("/publish", response_model=PublishResponse, status_code=status.HTTP_200_OK)
def publish(
    request: PublishRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office)
):
    """
    发起公示并直接分发到各教研室。
    - 同时将状态改为 distributed（教研室可以立即查看）
    - 记录操作日志
    """
    if not request.evaluation_ids:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="请至少选择一个教研室")

    evaluations = []
    for eval_id in request.evaluation_ids:
        evaluation = db.query(SelfEvaluation).filter(SelfEvaluation.id == eval_id).first()
        if not evaluation:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Evaluation {eval_id} not found"
            )
        if evaluation.status in ("published", "distributed"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"教研室 {evaluation.id} 的考评已经公示，不能重复公示"
            )
        if evaluation.status not in ("manually_scored", "ready_for_final", "finalized", "approved"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"教研室状态 '{evaluation.status}' 不满足公示条件，需要先完成评分"
            )
        evaluations.append(evaluation)

    now = datetime.utcnow()
    publication = Publication(
        evaluation_ids=[str(eval_id) for eval_id in request.evaluation_ids],
        published_by=current_user.id,
        published_at=now,
        distributed_at=now,  # 直接分发
    )
    db.add(publication)

    for evaluation in evaluations:
        evaluation.status = "distributed"

    db.flush()

    # 生成感悟总结
    for evaluation in evaluations:
        try:
            generate_insight_for_evaluation(db, evaluation.id)
        except Exception as e:
            logger.warning(f"Failed to generate insight for {evaluation.id}: {e}")

    # 操作日志
    operation_log = OperationLog(
        operation_type="publish",
        operator_id=current_user.id,
        operator_name=current_user.name,
        operator_role=current_user.role,
        target_id=publication.id,
        target_type="publication",
        details={
            "evaluation_ids": [str(eid) for eid in request.evaluation_ids],
            "evaluation_count": len(request.evaluation_ids),
            "action": "publish_and_distribute"
        }
    )
    db.add(operation_log)
    db.commit()
    db.refresh(publication)

    return PublishResponse(
        publication_id=publication.id,
        published_at=now,
        message=f"公示成功！{len(evaluations)} 个教研室的考评结果已分发，各教研室可立即查看结果。"
    )


@router.get("/publications", response_model=List[PublicationDetail])
def get_publications(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office)
):
    """查询公示记录列表"""
    publications = db.query(Publication).order_by(Publication.published_at.desc()).all()
    return [
        PublicationDetail(
            id=pub.id,
            evaluation_ids=[UUID(eid) for eid in pub.evaluation_ids],
            published_by=pub.published_by,
            published_at=pub.published_at,
            distributed_at=pub.distributed_at
        )
        for pub in publications
    ]


@router.get("/publications/{publication_id}", response_model=PublicationDetail)
def get_publication_detail(
    publication_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office)
):
    """查询单个公示记录详情"""
    publication = db.query(Publication).filter(Publication.id == publication_id).first()
    if not publication:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Publication not found")
    return PublicationDetail(
        id=publication.id,
        evaluation_ids=[UUID(eid) for eid in publication.evaluation_ids],
        published_by=publication.published_by,
        published_at=publication.published_at,
        distributed_at=publication.distributed_at
    )


@router.post("/distribute", response_model=DistributeResponse, status_code=status.HTTP_200_OK)
def distribute(
    request: DistributeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office)
):
    """分发结果到教研室端"""
    publication = db.query(Publication).filter(Publication.id == request.publication_id).first()
    if not publication:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Publication not found")
    if publication.distributed_at:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Already distributed at {publication.distributed_at}"
        )

    evaluation_ids = [UUID(eid) for eid in publication.evaluation_ids]
    evaluations = db.query(SelfEvaluation).filter(SelfEvaluation.id.in_(evaluation_ids)).all()

    distributed_at = datetime.utcnow()
    publication.distributed_at = distributed_at

    for evaluation in evaluations:
        evaluation.status = "distributed"

    for evaluation in evaluations:
        try:
            generate_insight_for_evaluation(db, evaluation.id)
        except Exception as e:
            logger.warning(f"Failed to generate insight for {evaluation.id}: {e}")

    operation_log = OperationLog(
        operation_type="distribute",
        operator_id=current_user.id,
        operator_name=current_user.name,
        operator_role=current_user.role,
        target_id=publication.id,
        target_type="publication",
        details={
            "publication_id": str(request.publication_id),
            "evaluation_ids": [str(eid) for eid in evaluation_ids],
            "evaluation_count": len(evaluations)
        }
    )
    db.add(operation_log)
    db.commit()
    db.refresh(publication)

    return DistributeResponse(
        distributed_count=len(evaluations),
        distributed_at=distributed_at,
        message=f"Results distributed successfully. {len(evaluations)} evaluation(s) accessible to teaching offices."
    )


@router.post("/sync-to-president")
def sync_to_president(
    db: Session = Depends(get_db),
    current_user: User = Depends(require_evaluation_office)
):
    """
    将已完成评分的考评信息上传至校长办公会端查看。
    返回所有 finalized/approved/published/distributed 状态的汇总数据。
    """
    evaluations = (
        db.query(SelfEvaluation)
        .filter(SelfEvaluation.status.in_(["finalized", "approved", "published", "distributed", "manually_scored"]))
        .all()
    )

    summary = []
    for ev in evaluations:
        office = ev.teaching_office
        final = db.query(FinalScore).filter(FinalScore.evaluation_id == ev.id).first()
        manual_scores = db.query(ManualScore).filter(ManualScore.evaluation_id == ev.id).all()
        manual_totals = [_manual_score_total(ms.scores) for ms in manual_scores]
        manual_avg = sum(manual_totals) / len(manual_totals) if manual_totals else None

        attachments = db.query(Attachment).filter(Attachment.evaluation_id == ev.id).all()

        summary.append({
            "id": str(ev.id),
            "teaching_office_name": office.name if office else "",
            "evaluation_year": ev.evaluation_year,
            "status": ev.status,
            "final_score": float(final.final_score) if final else None,
            "manual_score_avg": float(manual_avg) if manual_avg is not None else None,
            "manual_reviewer_count": len(manual_scores),
            "attachment_count": len(attachments),
            "submitted_at": ev.submitted_at.isoformat() if ev.submitted_at else None,
        })

    # 记录操作日志
    if evaluations:
        try:
            operation_log = OperationLog(
                operation_type="sync",
                operator_id=current_user.id,
                operator_name=current_user.name,
                operator_role=current_user.role,
                target_id=evaluations[0].id,
                target_type="sync_to_president",
                details={"evaluation_count": len(evaluations), "action": "sync_to_president"}
            )
            db.add(operation_log)
            db.commit()
        except Exception as e:
            logger.warning(f"Failed to log sync operation: {e}")

    return {
        "synced_count": len(summary),
        "synced_at": datetime.utcnow().isoformat(),
        "data": summary,
        "message": f"已将 {len(summary)} 条考评信息同步至校长办公会端"
    }


@router.get("/generate-result-word")
def generate_result_word(
    year: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    """
    生成考评结果Word文档，供校长端「同意发文」时下载。
    """
    # 查询已公示或已分发的考评
    query = (
        db.query(SelfEvaluation)
        .join(TeachingOffice, SelfEvaluation.teaching_office_id == TeachingOffice.id)
        .filter(SelfEvaluation.status.in_(["finalized", "approved", "published", "distributed"]))
    )
    if year:
        query = query.filter(SelfEvaluation.evaluation_year == year)
    evaluations = query.order_by(SelfEvaluation.submitted_at.desc()).all()

    # 汇总数据
    rows = []
    target_year = year or datetime.utcnow().year
    for ev in evaluations:
        office = ev.teaching_office
        final = db.query(FinalScore).filter(FinalScore.evaluation_id == ev.id).first()
        ai = db.query(AIScore).filter(AIScore.evaluation_id == ev.id).first()
        manual_scores = db.query(ManualScore).filter(ManualScore.evaluation_id == ev.id).all()
        manual_totals = [_manual_score_total(ms.scores) for ms in manual_scores]
        manual_avg = sum(manual_totals) / len(manual_totals) if manual_totals else None
        rows.append({
            "name": office.name if office else "",
            "year": ev.evaluation_year,
            "final_score": float(final.final_score) if final else None,
            "ai_score": float(ai.total_score) if ai else None,
            "manual_avg": float(manual_avg) if manual_avg is not None else None,
            "status": ev.status,
        })

    if not DOCX_AVAILABLE:
        # 若 python-docx 未安装，返回纯文本
        lines = [f"教研室工作考评结果汇总（{target_year}年度）\n"]
        lines.append(f"{'教研室名称':<20}{'年度':<8}{'最终得分':<12}{'AI评分':<12}{'人工均分':<12}\n")
        lines.append("-" * 64 + "\n")
        for r in rows:
            fs = f"{r['final_score']:.1f}" if r['final_score'] is not None else "-"
            ai = f"{r['ai_score']:.1f}" if r['ai_score'] is not None else "-"
            ma = f"{r['manual_avg']:.1f}" if r['manual_avg'] is not None else "-"
            lines.append(f"{r['name']:<20}{r['year']:<8}{fs:<12}{ai:<12}{ma:<12}\n")
        content = "".join(lines).encode("utf-8")
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="result_{target_year}.txt"'}
        )

    # 生成 Word 文档
    doc = DocxDocument()
    doc.add_heading(f"教研室工作考评结果汇总（{target_year}年度）", level=0)
    doc.add_paragraph(f"发文时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"共 {len(rows)} 个教研室")
    doc.add_paragraph("")

    # 表格
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["教研室名称", "考评年度", "最终得分", "AI评分", "人工均分"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # 排序：最终得分降序
    rows_sorted = sorted(rows, key=lambda x: x["final_score"] or 0, reverse=True)
    for r in rows_sorted:
        row_cells = table.add_row().cells
        row_cells[0].text = r["name"]
        row_cells[1].text = str(r["year"])
        row_cells[2].text = f"{r['final_score']:.1f}" if r["final_score"] is not None else "-"
        row_cells[3].text = f"{r['ai_score']:.1f}" if r["ai_score"] is not None else "-"
        row_cells[4].text = f"{r['manual_avg']:.1f}" if r["manual_avg"] is not None else "-"

    doc.add_paragraph("")
    doc.add_paragraph("备注：本文档由系统自动生成，仅供参考。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"result_{target_year}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

